import pdfplumber
from docx import Document

def load_pdf(path):
    texts = []

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            # Extract normal text
            if page.extract_text():
                texts.append(page.extract_text())

            # Extract tables
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    clean_row = " | ".join(cell or "" for cell in row)
                    texts.append(clean_row)

    return "\n".join(texts)


def load_docx(path):
    doc = Document(path)
    texts = []

    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            texts.append(
                " | ".join(cell.text.strip() for cell in row.cells)
            )

    return "\n".join(texts)


def load_txt(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def load_document(path):
    if path.endswith(".pdf"):
        return load_pdf(path)
    if path.endswith(".docx"):
        return load_docx(path)
    if path.endswith(".txt"):
        return load_txt(path)
    if path.endswith(".puml"):
        return load_txt(path)
    raise ValueError("Unsupported file type")
